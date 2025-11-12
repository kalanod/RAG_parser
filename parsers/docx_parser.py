from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import List, Optional, Tuple

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from docx import Document as DocxDocument
from docx.document import Document as _DocxDocumentType
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from PIL import Image
import pytesseract
import logging

from parsers.abstract_document_parser import AbstractDocumentParser

LOGGER = logging.getLogger(__name__)


@dataclass
class DocxParserConfig:
    chunk_size: int = 700
    chunk_overlap: int = 150
    overlap_chars: int = 300
    ocr_lang: str = "rus+eng"
    min_ocr_w: int = 100
    min_ocr_h: int = 50


class DOCXParser(AbstractDocumentParser):
    supported_extensions = (".docx",)
    config = DocxParserConfig()

    def parse(self, path: Path) -> List[Document]:
        try:
            documents: List[Document] = []
            documents.extend(self._extract_text_blocks(path))
            documents.extend(self._extract_table_blocks(path))
            documents.extend(self._extract_ocr_blocks(path))
            return self.split_to_docs(documents)
        except Exception as exc:
            LOGGER.exception("Failed to parse DOCX %s: %s", path, exc)
            return [
                Document(
                    page_content=f"[ERROR] Failed to load document: {str(exc)}",
                    metadata={
                        "source": str(path),
                        "error": True,
                        "error_type": type(exc).__name__,
                    },
                )
            ]

    def split_to_docs(self, documents: List[Document]) -> List[Document]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            separators=["\n\n", "\n", ".", "?", "!", " "],
        )
        split_docs: List[Document] = []

        for doc in documents:
            chunks = splitter.split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                split_docs.append(
                    Document(
                        page_content=chunk,
                        metadata={
                            **doc.metadata,
                            "chunk_index": i,
                            "chunk_count": len(chunks),
                        },
                    )
                )
        return split_docs

    def _extract_text_blocks(self, path: Path) -> List[Document]:
        d: _DocxDocumentType = DocxDocument(str(path))
        paragraphs: List[Paragraph] = [p for p in d.paragraphs if p.text and p.text.strip()]

        lines: List[Tuple[int, str, Optional[str]]] = []
        for idx, p in enumerate(paragraphs):
            text = p.text.strip()
            style_name = getattr(p.style, "name", None)
            if text:
                lines.append((idx, text, style_name))

        docs: List[Document] = []
        N = len(lines)
        for i in range(N):
            idx, cur, style = lines[i]
            nxt = lines[i + 1][1] if i + 1 < N else ""
            block = cur + ("\n" + nxt[: self.config.overlap_chars] if nxt else "")

            if block.strip():
                docs.append(
                    Document(
                        page_content=block.strip(),
                        metadata={
                            "source": str(path),
                            "type": "paragraph_overlap",
                            "paragraph_index": idx,
                            "style": style or "",
                            "overlap_with_next": bool(nxt),
                        },
                    )
                )

        return docs

    def _extract_table_blocks(self, path: Path) -> List[Document]:
        docs: List[Document] = []
        d: _DocxDocumentType = DocxDocument(str(path))

        def cell_text(cell: _Cell) -> str:
            parts = [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
            return " ".join(parts).strip()

        for t_idx, table in enumerate(d.tables):
            if not isinstance(table, Table):
                continue
            rows_text: List[str] = []
            for row in table.rows:
                cells = [cell_text(c) for c in row.cells]
                row_line = " | ".join(c or "" for c in cells)
                if any(cells):
                    rows_text.append(row_line)

            table_text = "\n".join(rows_text).strip()
            if table_text:
                docs.append(
                    Document(
                        page_content=table_text,
                        metadata={
                            "source": str(path),
                            "type": "table",
                            "table_index": t_idx,
                        },
                    )
                )

        return docs

    def _extract_ocr_blocks(self, path: Path) -> List[Document]:
        docs: List[Document] = []
        d: _DocxDocumentType = DocxDocument(str(path))
        for r_id, part in d.part.related_parts.items():
            content_type: str = getattr(part, "content_type", "")
            if not content_type.startswith("image/"):
                continue

            try:
                blob: bytes = part.blob  # raw image bytes
                img = Image.open(BytesIO(blob))
                w, h = img.size
                if w < self.config.min_ocr_w or h < self.config.min_ocr_h:
                    continue

                try:
                    ocr_text = pytesseract.image_to_string(img, lang=self.config.ocr_lang).strip()
                except Exception as e:
                    LOGGER.warning("OCR failed for image %s in %s: %s", r_id, path, e)
                    ocr_text = ""

                if ocr_text:
                    docs.append(
                        Document(
                            page_content=ocr_text,
                            metadata={
                                "source": str(path),
                                "type": "image_text",
                                "relation_id": r_id,
                                "image_size": f"{w}x{h}",
                            },
                        )
                    )
            except Exception as e:
                LOGGER.debug("Skipping image %s due to error: %s", r_id, e)
                continue

        return docs
